"""
ChromaDB Knowledge Base Service
Refactored document/chunk indexing service with automatic metadata generation.
"""
import os
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

import chromadb
from chromadb.config import Settings
from fastapi import BackgroundTasks, Body, FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field

import config
from kb_logic import (
    build_document_metadata,
    build_vector_filter,
    build_vector_metadata,
    chunk_document,
    chunk_matches_session,
    classify_chunks,
    clean_html,
    compute_keyword_overlap,
    default_title_from_filename,
    detect_language,
    infer_source_org,
    map_query_topics,
    normalize_title_key,
    now_iso,
    sha256_bytes,
    summarize_text,
    tokenize_query,
)
from kb_repository import KnowledgeRepository


ALLOWED_CHILD_AGE_BANDS = {"0_3", "3_6", "6_12", "12_18"}
ALLOWED_AGE_BANDS = {"all", *ALLOWED_CHILD_AGE_BANDS}
ALLOWED_DOC_LIBRARIES = {"rules", "parent", "age_content", "mixed"}
ALLOWED_DOC_AUDIENCES = {"system", "parent", "child", "teacher", "mixed"}
ALLOWED_CHUNK_AUDIENCES = {"system", "parent", "child", "teacher"}
ALLOWED_DOC_VISIBILITIES = {"system_only", "parent_only", "retrieval_visible", "mixed"}
ALLOWED_CHUNK_VISIBILITIES = {"system_only", "parent_only", "retrieval_visible", "blocked"}
ALLOWED_REVIEW_STATUS = {"auto_accepted", "needs_review", "blocked"}
ALLOWED_RISK_LEVEL = {"low", "medium", "high"}
ALLOWED_ACTIONS = {"reanalyze", "reparse", "rechunk", "reindex"}
MAX_RETRIEVAL_CANDIDATES = 100
RELEVANCE_SCORE_THRESHOLD = 0.52
DISTANCE_THRESHOLD = 0.85
STRONG_SEMANTIC_DISTANCE = 0.45


app = FastAPI(
    title="Knowledge Base Service",
    description="ChromaDB-powered knowledge base for AI Chat",
    version="2.1.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs(config.UPLOAD_DIR, exist_ok=True)
os.makedirs(config.CHROMA_PERSIST_DIR, exist_ok=True)

chroma_client = chromadb.PersistentClient(
    path=config.CHROMA_PERSIST_DIR,
    settings=Settings(anonymized_telemetry=False),
)

collection = chroma_client.get_or_create_collection(
    name=config.CHROMA_COLLECTION_NAME,
    metadata={"description": "Knowledge base chunk collection"},
)

_repository: Optional[KnowledgeRepository] = None


def get_repository() -> KnowledgeRepository:
    global _repository
    if _repository is None:
        _repository = KnowledgeRepository()
    return _repository


class FileInfo(BaseModel):
    id: str
    title: str
    filename: str
    original_filename: str
    file_type: str
    file_size: int
    upload_time: str
    updated_at: str = ""
    chunk_count: int
    version: int = 1
    language: str = "other"
    source_org: str = "Unknown"
    library: str = "mixed"
    audience: str = "mixed"
    age_bands: List[str] = Field(default_factory=list)
    safety_visibility: str = "mixed"
    topics: List[str] = Field(default_factory=list)
    summary: str = ""
    risk_level: str = "low"
    review_status: str = "needs_review"
    parser_status: str = "pending"
    chunk_status: str = "pending"
    embedding_status: str = "pending"
    indexing_status: str = "pending"
    enabled: bool = True
    last_indexed_at: Optional[str] = None
    error_message: Optional[str] = None


class UploadResponse(BaseModel):
    success: bool
    message: str
    duplicate: bool = False
    processing: bool = False
    file: Optional[FileInfo] = None


class FileListResponse(BaseModel):
    success: bool
    files: List[FileInfo]
    total: int


class SearchResult(BaseModel):
    document: str
    metadata: Dict[str, Any]
    distance: float
    score: float = 0.0
    reliable: bool = False
    passed_relevance_threshold: bool = False
    reason: str = ""
    match_signals: List[str] = Field(default_factory=list)


class SearchResponse(BaseModel):
    success: bool
    results: List[SearchResult]
    filtered_out: List[Dict[str, Any]] = Field(default_factory=list)
    reliable: bool = False
    message: str = ""
    no_result_reason: Optional[str] = None
    query_topics: List[str] = Field(default_factory=list)
    topic_filter_applied: bool = False
    missing_topic_content: bool = False
    query: str
    session_type: str
    age_band: Optional[str] = None
    indexed_document_count: int = 0


class DeleteResponse(BaseModel):
    success: bool
    message: str


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def validate_filename(filename: str) -> None:
    if len(filename) > config.MAX_FILENAME_LENGTH:
        raise HTTPException(
            status_code=400,
            detail="Filename too long. Maximum {0} characters allowed.".format(config.MAX_FILENAME_LENGTH),
        )

    ext = get_file_extension(filename)
    if ext not in config.ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="File type not allowed. Allowed types: {0}".format(", ".join(sorted(config.ALLOWED_EXTENSIONS))),
        )


def validate_title(title: str) -> str:
    normalized = title.strip()
    if not normalized:
        raise HTTPException(status_code=400, detail="Title cannot be empty")
    if len(normalized) > 255:
        raise HTTPException(status_code=400, detail="Title too long. Maximum 255 characters allowed.")
    return normalized


def read_text_file(file_path: str) -> str:
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as handle:
                return handle.read()
        except UnicodeDecodeError:
            continue
    return ""


def extract_text_from_file(file_path: str, file_format: str) -> str:
    if file_format in {"txt", "md"}:
        return read_text_file(file_path).strip()

    if file_format == "html":
        return clean_html(read_text_file(file_path))

    if file_format == "pdf":
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise HTTPException(status_code=500, detail="PyPDF2 is required for PDF parsing") from exc

        reader = PdfReader(file_path)
        parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
        return "\n".join(parts).strip()

    if file_format == "docx":
        try:
            import docx
        except ImportError as exc:
            raise HTTPException(status_code=500, detail="python-docx is required for DOCX parsing") from exc

        document = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in document.paragraphs]).strip()

    raise HTTPException(status_code=400, detail="Unsupported file format for parsing")


def ordered_age_bands(values: List[str]) -> List[str]:
    order = ["all", "0_3", "3_6", "6_12", "12_18"]
    unique = []
    for value in values:
        if value not in unique:
            unique.append(value)
    return [item for item in order if item in unique]


def normalize_list_value(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        return [item.strip() for item in value.split(",") if item.strip()]
    return []


def normalize_age_bands(value: Any) -> List[str]:
    age_bands = normalize_list_value(value)
    invalid = [item for item in age_bands if item not in ALLOWED_AGE_BANDS]
    if invalid:
        raise HTTPException(status_code=400, detail="Invalid age_bands value")
    return ordered_age_bands(age_bands)


def build_failed_document_metadata(
    doc_id: str,
    title: str,
    original_filename: str,
    file_ext: str,
    mime_type: str,
    file_format: str,
    storage_key: str,
    content_hash: str,
    version: int,
    error_message: str,
) -> Dict[str, Any]:
    now = now_iso()
    return {
        "doc_id": doc_id,
        "document_key": normalize_title_key(title),
        "title": title,
        "original_filename": original_filename,
        "file_ext": file_ext,
        "mime_type": mime_type,
        "format": file_format,
        "storage_key": storage_key,
        "content_hash": content_hash,
        "version": version,
        "language": "other",
        "source_org": "Unknown",
        "source_org_confidence": 0.2,
        "library": "mixed",
        "library_confidence": 0.2,
        "audience": "mixed",
        "audience_confidence": 0.2,
        "age_bands": ["all"],
        "age_bands_confidence": 0.2,
        "safety_visibility": "mixed",
        "topics": [],
        "tags": ["parse_failed"],
        "summary": error_message,
        "error_message": error_message,
        "risk_level": "high",
        "enabled": False,
        "review_status": "blocked",
        "parser_status": config.STATUS_FAILED,
        "chunk_status": config.STATUS_FAILED,
        "embedding_status": config.STATUS_FAILED,
        "indexing_status": config.STATUS_FAILED,
        "created_at": now,
        "updated_at": now,
        "last_indexed_at": None,
    }


def build_pending_document_metadata(
    doc_id: str,
    title: str,
    original_filename: str,
    file_ext: str,
    mime_type: str,
    file_format: str,
    storage_key: str,
    content_hash: str,
    version: int,
) -> Dict[str, Any]:
    now = now_iso()
    return {
        "doc_id": doc_id,
        "document_key": normalize_title_key(title),
        "title": title,
        "original_filename": original_filename,
        "file_ext": file_ext,
        "mime_type": mime_type,
        "format": file_format,
        "storage_key": storage_key,
        "content_hash": content_hash,
        "version": version,
        "language": "other",
        "source_org": "Unknown",
        "source_org_confidence": 0.0,
        "library": "mixed",
        "library_confidence": 0.0,
        "audience": "mixed",
        "audience_confidence": 0.0,
        "age_bands": [],
        "age_bands_confidence": 0.0,
        "safety_visibility": "mixed",
        "topics": [],
        "tags": [],
        "summary": "Queued for automatic analysis.",
        "error_message": None,
        "risk_level": "low",
        "enabled": True,
        "review_status": "needs_review",
        "parser_status": config.STATUS_PENDING,
        "chunk_status": config.STATUS_PENDING,
        "embedding_status": config.STATUS_PENDING,
        "indexing_status": config.STATUS_PENDING,
        "created_at": now,
        "updated_at": now,
        "last_indexed_at": None,
    }


def to_file_info(document: Dict[str, Any], chunk_count: int = 0) -> FileInfo:
    filename = os.path.basename(str(document.get("storage_key") or ""))
    return FileInfo(
        id=str(document.get("doc_id") or ""),
        title=str(document.get("title") or ""),
        filename=filename,
        original_filename=str(document.get("original_filename") or ""),
        file_type=str(document.get("file_ext") or ""),
        file_size=int(document.get("file_size") or 0),
        upload_time=str(document.get("created_at") or document.get("updated_at") or ""),
        updated_at=str(document.get("updated_at") or ""),
        chunk_count=chunk_count or int(document.get("chunk_count") or 0),
        version=int(document.get("version") or 1),
        language=str(document.get("language") or "other"),
        source_org=str(document.get("source_org") or "Unknown"),
        library=str(document.get("library") or "mixed"),
        audience=str(document.get("audience") or "mixed"),
        age_bands=list(document.get("age_bands") or []),
        safety_visibility=str(document.get("safety_visibility") or "mixed"),
        topics=list(document.get("topics") or []),
        summary=str(document.get("summary") or ""),
        risk_level=str(document.get("risk_level") or "low"),
        review_status=str(document.get("review_status") or "needs_review"),
        parser_status=str(document.get("parser_status") or config.STATUS_PENDING),
        chunk_status=str(document.get("chunk_status") or config.STATUS_PENDING),
        embedding_status=str(document.get("embedding_status") or config.STATUS_PENDING),
        indexing_status=str(document.get("indexing_status") or config.STATUS_PENDING),
        enabled=bool(document.get("enabled")),
        last_indexed_at=document.get("last_indexed_at"),
        error_message=str(document.get("error_message") or "") or None,
    )


def get_storage_path(storage_key: str) -> str:
    return os.path.join(config.STORAGE_DIR, storage_key)


def get_status_label(status: str) -> str:
    if status == config.STATUS_COMPLETED:
        return "completed"
    if status == config.STATUS_FAILED:
        return "failed"
    return "pending"


def build_processing_stages(document: Dict[str, Any]) -> List[Dict[str, str]]:
    parser_status = str(document.get("parser_status") or config.STATUS_PENDING)
    chunk_status = str(document.get("chunk_status") or config.STATUS_PENDING)
    embedding_status = str(document.get("embedding_status") or config.STATUS_PENDING)
    indexing_status = str(document.get("indexing_status") or config.STATUS_PENDING)

    classification_status = "failed" if parser_status == config.STATUS_FAILED else (
        "completed" if parser_status == config.STATUS_COMPLETED else "pending"
    )
    metadata_status = "failed" if chunk_status == config.STATUS_FAILED else (
        "completed" if chunk_status == config.STATUS_COMPLETED else "pending"
    )
    embedding_stage_status = "failed" if embedding_status == config.STATUS_FAILED or indexing_status == config.STATUS_FAILED else (
        "completed" if embedding_status == config.STATUS_COMPLETED and indexing_status == config.STATUS_COMPLETED else "pending"
    )
    final_status = "failed" if indexing_status == config.STATUS_FAILED else (
        "completed" if indexing_status == config.STATUS_COMPLETED else "pending"
    )

    return [
        {"key": "saved", "label": "原文件已保存", "status": "completed"},
        {"key": "parser", "label": "文档解析中", "status": get_status_label(parser_status)},
        {"key": "chunking", "label": "文本切片中", "status": get_status_label(chunk_status)},
        {"key": "classification", "label": "自动分类中", "status": classification_status},
        {"key": "metadata", "label": "自动生成年龄段/受众/可见性/主题中", "status": metadata_status},
        {"key": "embedding", "label": "构建 embedding / 建立索引中", "status": embedding_stage_status},
        {"key": "done", "label": "分析完成", "status": final_status},
    ]


def get_document_status_payload(document: Dict[str, Any], chunk_count: int = 0) -> Dict[str, Any]:
    file_info = to_file_info(document, chunk_count=chunk_count)
    stages = build_processing_stages(document)
    completed = str(document.get("indexing_status") or "") == config.STATUS_COMPLETED
    failed = any(stage["status"] == "failed" for stage in stages)

    current_stage = "saved"
    for stage in stages:
        current_stage = stage["key"]
        if stage["status"] == "pending":
            break

    return {
        "success": True,
        "file": file_info.model_dump(),
        "stages": stages,
        "completed": completed,
        "failed": failed,
        "current_stage": current_stage,
        "error_message": str(document.get("error_message") or "") or None,
    }


def to_chunk_payload(chunk: Dict[str, Any]) -> Dict[str, Any]:
    heading_path = chunk.get("heading_path") or []
    if isinstance(heading_path, list):
        heading_display = " / ".join([str(item) for item in heading_path if str(item).strip()])
    else:
        heading_display = str(heading_path)

    return {
        "chunk_id": str(chunk.get("chunk_id") or ""),
        "doc_id": str(chunk.get("doc_id") or ""),
        "chunk_index": int(chunk.get("chunk_index") or 0),
        "heading_path": heading_path,
        "heading_path_display": heading_display or "Root",
        "content": str(chunk.get("content") or ""),
        "content_preview": str(chunk.get("content") or "")[:240],
        "char_count": int(chunk.get("char_count") or 0),
        "token_count": int(chunk.get("token_count") or 0),
        "age_bands": list(chunk.get("age_bands") or []),
        "audience": str(chunk.get("audience") or "child"),
        "visibility": str(chunk.get("visibility") or "blocked"),
        "topics": list(chunk.get("topics") or []),
        "risk_level": str(chunk.get("risk_level") or "low"),
        "retrieval_enabled": bool(chunk.get("retrieval_enabled")),
        "confidence": float(chunk.get("confidence") or 0.0),
        "chunk_summary": str(chunk.get("chunk_summary") or ""),
        "embedding_id": chunk.get("embedding_id"),
        "vector_id": chunk.get("vector_id"),
        "created_at": str(chunk.get("created_at") or ""),
        "updated_at": str(chunk.get("updated_at") or ""),
    }


def explain_filter_failure(metadata: Dict[str, Any], session_type: str, age_band: Optional[str]) -> List[str]:
    reasons = []
    if int(metadata.get("enabled", 0)) != 1:
        reasons.append("document_disabled")
    if int(metadata.get("retrieval_enabled", 0)) != 1:
        reasons.append("chunk_retrieval_disabled")

    visibility = str(metadata.get("visibility", "blocked"))
    if session_type == "child":
        if visibility != "retrieval_visible":
            reasons.append("visibility_not_child_safe")
        if age_band in ALLOWED_CHILD_AGE_BANDS:
            if int(metadata.get("age_all", 0)) != 1 and int(metadata.get("age_{0}".format(age_band), 0)) != 1:
                reasons.append("age_band_mismatch")
    elif session_type == "parent":
        if visibility not in {"parent_only", "retrieval_visible"}:
            reasons.append("visibility_not_parent_allowed")
    elif session_type == "system":
        if visibility not in {"system_only", "parent_only", "retrieval_visible"}:
            reasons.append("visibility_not_system_allowed")
    else:
        reasons.append("invalid_session_type")

    return reasons


def build_result_reason(metadata: Dict[str, Any], session_type: str, age_band: Optional[str], match_signals: Optional[List[str]] = None) -> str:
    visibility = str(metadata.get("visibility", "blocked"))
    audience = str(metadata.get("audience", "child"))
    reasons = list(match_signals or [])

    if session_type == "child":
        if age_band in ALLOWED_CHILD_AGE_BANDS:
            reasons.append("passed age filter {0}".format(age_band))
        else:
            reasons.append("no age filter applied")
    elif session_type == "parent":
        reasons.append("passed parent visibility filter")
    else:
        reasons.append("passed system visibility filter")

    reasons.append("visibility={0}".format(visibility))
    reasons.append("audience={0}".format(audience))
    return ", ".join(reasons)


def merge_where_filters(base_filter: Dict[str, Any], extra_conditions: List[Dict[str, Any]]) -> Dict[str, Any]:
    base_conditions = list(base_filter.get("$and", []))
    return {"$and": base_conditions + extra_conditions}


def build_topic_filter(query_topics: List[str]) -> Optional[Dict[str, Any]]:
    if not query_topics:
        return None

    conditions = []
    for topic in query_topics:
        conditions.append({"topic_{0}".format(topic): 1})

    if not conditions:
        return None
    if len(conditions) == 1:
        return conditions[0]
    return {"$or": conditions}


def run_vector_query(query: str, n_results: int, where_filter: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    try:
        if where_filter:
            return collection.query(query_texts=[query], n_results=n_results, where=where_filter)
        return collection.query(query_texts=[query], n_results=n_results)
    except Exception:
        return collection.query(query_texts=[query], n_results=n_results)


def hydrate_search_metadata(raw_metadata: Dict[str, Any], cache: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
    repository = get_repository()
    metadata = dict(raw_metadata or {})
    doc_id = str(metadata.get("doc_id") or "")
    chunk_id = str(metadata.get("chunk_id") or "")

    if doc_id:
        if doc_id not in cache["documents"]:
            cache["documents"][doc_id] = repository.get_document(doc_id) or {}
        document = cache["documents"][doc_id]
        metadata["title"] = str(document.get("title") or metadata.get("title") or "")
        metadata["original_filename"] = str(document.get("original_filename") or metadata.get("original_filename") or "")
        metadata["source_org"] = str(document.get("source_org") or metadata.get("source_org") or "Unknown")
        metadata["language"] = str(document.get("language") or metadata.get("language") or "other")
        metadata["review_status"] = str(document.get("review_status") or metadata.get("review_status") or "needs_review")
        metadata["document_topics"] = list(document.get("topics") or [])

    if chunk_id:
        if chunk_id not in cache["chunks"]:
            cache["chunks"][chunk_id] = repository.get_chunk(chunk_id) or {}
        chunk = cache["chunks"][chunk_id]
        heading_path = chunk.get("heading_path") or []
        if isinstance(heading_path, list):
            metadata["heading_path"] = heading_path
            metadata["heading_path_display"] = " / ".join([str(item) for item in heading_path if str(item).strip()]) or "Root"
        else:
            metadata["heading_path"] = heading_path
            metadata["heading_path_display"] = str(heading_path or "Root")
        metadata["topics"] = list(chunk.get("topics") or [])
        metadata["age_bands"] = list(chunk.get("age_bands") or [])
        metadata["visibility"] = str(chunk.get("visibility") or metadata.get("visibility") or "blocked")
        metadata["audience"] = str(chunk.get("audience") or metadata.get("audience") or "child")
        metadata["retrieval_enabled"] = bool(chunk.get("retrieval_enabled")) if chunk else bool(metadata.get("retrieval_enabled"))
        metadata["confidence"] = float(chunk.get("confidence") or metadata.get("confidence") or 0.0)

    return metadata


def compute_semantic_score(distance: float) -> float:
    bounded = max(0.0, min(float(distance), 1.5))
    return round(max(0.0, 1.0 - (bounded / 1.5)), 4)


def evaluate_candidate_relevance(
    query: str,
    query_terms: List[str],
    query_topics: List[str],
    document: str,
    metadata: Dict[str, Any],
    distance: float,
    session_type: str,
    age_band: Optional[str],
) -> Dict[str, Any]:
    title = str(metadata.get("title") or "")
    heading = str(metadata.get("heading_path_display") or metadata.get("heading_path") or "")
    topics = list(metadata.get("topics") or [])
    age_bands = list(metadata.get("age_bands") or [])

    semantic_score = compute_semantic_score(distance)
    keyword_overlap = compute_keyword_overlap(query_terms, [document, title, heading, " ".join(topics)])
    title_hit = compute_keyword_overlap(query_terms, [title]) > 0.0
    heading_hit = compute_keyword_overlap(query_terms, [heading]) > 0.0
    topic_hit = bool(query_topics and any(topic in topics for topic in query_topics))
    document_topic_hit = bool(query_topics and any(topic in list(metadata.get("document_topics") or []) for topic in query_topics))

    score = round(min(1.0,
        (semantic_score * 0.45)
        + (keyword_overlap * 0.30)
        + (0.20 if topic_hit else 0.0)
        + (0.10 if title_hit or heading_hit else 0.0)
        + (0.05 if document_topic_hit else 0.0),
    ), 4)

    match_signals = []
    if topic_hit:
        match_signals.append("topic match")
    elif document_topic_hit:
        match_signals.append("document topic match")
    if title_hit:
        match_signals.append("title keyword match")
    if heading_hit:
        match_signals.append("heading match")
    if keyword_overlap >= 0.34:
        match_signals.append("keyword overlap")
    if distance <= STRONG_SEMANTIC_DISTANCE:
        match_signals.append("semantic match")

    rejection_reasons = []
    if distance > DISTANCE_THRESHOLD:
        rejection_reasons.append("distance_above_threshold")
    if query_topics and not topic_hit and not document_topic_hit:
        rejection_reasons.append("topic_mismatch")
    if keyword_overlap <= 0 and not title_hit and not heading_hit and not topic_hit and distance > STRONG_SEMANTIC_DISTANCE:
        rejection_reasons.append("no_keyword_or_topic_support")
    if score < RELEVANCE_SCORE_THRESHOLD:
        rejection_reasons.append("below_relevance_threshold")

    reliable = not rejection_reasons and (
        topic_hit
        or keyword_overlap >= 0.34
        or title_hit
        or heading_hit
        or distance <= STRONG_SEMANTIC_DISTANCE
    )

    return {
        "score": score,
        "semantic_score": semantic_score,
        "keyword_overlap": keyword_overlap,
        "topic_hit": topic_hit,
        "document_topic_hit": document_topic_hit,
        "title_hit": title_hit,
        "heading_hit": heading_hit,
        "match_signals": match_signals,
        "reliable": reliable,
        "rejection_reasons": rejection_reasons,
        "passed_relevance_threshold": reliable,
        "reason": build_result_reason(metadata, session_type, age_band, match_signals),
        "age_bands": age_bands,
    }


def build_search_message(reliable: bool, query_topics: List[str], missing_topic_content: bool, results_count: int) -> Dict[str, Any]:
    if missing_topic_content and query_topics:
        return {
            "reliable": False,
            "message": "当前知识库缺少该主题内容",
            "no_result_reason": "current_topic_missing",
        }
    if reliable and results_count > 0:
        return {
            "reliable": True,
            "message": "已返回通过相关性阈值的真实检索结果",
            "no_result_reason": None,
        }
    return {
        "reliable": False,
        "message": "当前知识库中没有找到可靠的匹配结果",
        "no_result_reason": "no_reliable_match",
    }


def query_collection_with_diagnostics(query: str, session_type: str, age_band: Optional[str], limit: int) -> Dict[str, Any]:
    n_results = min(max(limit * 12, 30), MAX_RETRIEVAL_CANDIDATES)
    query_topics = map_query_topics(query)
    topic_filter = build_topic_filter(query_topics)
    topic_filter_applied = bool(topic_filter)
    raw = None
    missing_topic_content = False

    if topic_filter:
        raw = run_vector_query(query, n_results, topic_filter)
        topic_docs = (raw.get("documents") or [[]])[0] if raw else []
        if not topic_docs:
            missing_topic_content = True
            raw = {"documents": [[]], "metadatas": [[]], "distances": [[]]}
    else:
        raw = run_vector_query(query, n_results)

    documents = raw.get("documents", [[]])
    metadatas = raw.get("metadatas", [[]])
    distances = raw.get("distances", [[]])

    docs = documents[0] if documents else []
    meta_items = metadatas[0] if metadatas else []
    distance_items = distances[0] if distances else []

    allowed_results: List[SearchResult] = []
    filtered_out: List[Dict[str, Any]] = []
    cache = {"documents": {}, "chunks": {}}
    query_terms = tokenize_query(query)

    for index, document in enumerate(docs):
        raw_metadata = meta_items[index] if index < len(meta_items) else {}
        metadata = hydrate_search_metadata(raw_metadata, cache)
        distance = distance_items[index] if index < len(distance_items) else 0.0
        reasons = explain_filter_failure(metadata, session_type, age_band)

        if reasons:
            filtered_out.append(
                {
                    "document": document,
                    "metadata": metadata,
                    "distance": distance,
                    "score": 0.0,
                    "reliable": False,
                    "passed_relevance_threshold": False,
                    "reasons": reasons,
                }
            )
            continue

        relevance = evaluate_candidate_relevance(
            query=query,
            query_terms=query_terms,
            query_topics=query_topics,
            document=document,
            metadata=metadata,
            distance=distance,
            session_type=session_type,
            age_band=age_band,
        )

        if relevance["reliable"] and len(allowed_results) < limit:
            allowed_results.append(
                SearchResult(
                    document=document,
                    metadata=metadata,
                    distance=distance,
                    score=relevance["score"],
                    reliable=True,
                    passed_relevance_threshold=True,
                    reason=relevance["reason"],
                    match_signals=relevance["match_signals"],
                )
            )
            continue

        filtered_out.append(
            {
                "document": document,
                "metadata": metadata,
                "distance": distance,
                "score": relevance["score"],
                "reliable": False,
                "passed_relevance_threshold": False,
                "reasons": relevance["rejection_reasons"] or ["below_relevance_threshold"],
                "match_signals": relevance["match_signals"],
            }
        )

    message_info = build_search_message(bool(allowed_results), query_topics, missing_topic_content, len(allowed_results))
    return {
        "results": allowed_results,
        "filtered_out": filtered_out,
        "query_topics": query_topics,
        "topic_filter_applied": topic_filter_applied,
        "missing_topic_content": missing_topic_content,
        "reliable": message_info["reliable"],
        "message": message_info["message"],
        "no_result_reason": message_info["no_result_reason"],
    }


def query_collection(query: str, session_type: str, age_band: Optional[str], limit: int) -> List[SearchResult]:
    diagnostics = query_collection_with_diagnostics(query, session_type, age_band, limit)
    return diagnostics["results"]


def update_vector_titles(doc_id: str, title: str) -> None:
    repository = get_repository()
    chunks = repository.get_chunks_for_doc(doc_id)
    chunk_ids = [chunk["vector_id"] for chunk in chunks if chunk.get("vector_id")]
    if not chunk_ids:
        return

    vector_rows = collection.get(ids=chunk_ids, include=["metadatas"])
    ids = vector_rows.get("ids", [])
    metadatas = vector_rows.get("metadatas", [])
    if not ids or not metadatas:
        return

    updated = []
    for metadata in metadatas:
        new_metadata = dict(metadata or {})
        new_metadata["title"] = title
        updated.append(new_metadata)

    collection.update(ids=ids, metadatas=updated)


def sync_document_vectors(doc_id: str, chunk_ids: Optional[List[str]] = None) -> None:
    repository = get_repository()
    document = repository.get_document(doc_id)
    if not document:
        return

    chunks = repository.get_chunks_for_doc(doc_id)
    if chunk_ids:
        wanted = set(chunk_ids)
        chunks = [chunk for chunk in chunks if str(chunk.get("chunk_id")) in wanted]

    vector_ids = [str(chunk.get("vector_id")) for chunk in chunks if chunk.get("vector_id")]
    if not vector_ids:
        return

    collection.upsert(
        ids=vector_ids,
        documents=[str(chunk.get("content") or "") for chunk in chunks],
        metadatas=[build_vector_metadata(document, chunk) for chunk in chunks],
    )


class KnowledgeIngestionPipeline:
    def __init__(self, repository: KnowledgeRepository):
        self.repository = repository

    def saveOriginalFile(self, doc_id: str, content: bytes, file_ext: str, version: int) -> str:
        safe_filename = "{0}_v{1}{2}".format(doc_id, version, file_ext)
        target_path = os.path.join(config.UPLOAD_DIR, safe_filename)
        with open(target_path, "wb") as handle:
            handle.write(content)
        return os.path.join("uploads", safe_filename)

    def parseDocument(self, storage_key: str, file_format: str) -> str:
        file_path = os.path.join(config.STORAGE_DIR, storage_key)
        return extract_text_from_file(file_path, file_format)

    def detectLanguage(self, text: str) -> str:
        return detect_language(text)

    def inferSourceOrg(self, title: str, original_filename: str, text: str):
        return infer_source_org(title, original_filename, text)

    def classifyDocument(
        self,
        doc_id: str,
        title: str,
        original_filename: str,
        file_ext: str,
        mime_type: str,
        file_format: str,
        storage_key: str,
        content_hash: str,
        version: int,
        text: str,
    ) -> Dict[str, Any]:
        return build_document_metadata(
            doc_id=doc_id,
            title=title,
            original_filename=original_filename,
            file_ext=file_ext,
            mime_type=mime_type,
            file_format=file_format,
            storage_key=storage_key,
            content_hash=content_hash,
            version=version,
            text=text,
        )

    def chunkDocument(self, text: str) -> List[Dict[str, Any]]:
        return chunk_document(text, config.CHUNK_SIZE, config.CHUNK_OVERLAP)

    def classifyChunks(self, chunks: List[Dict[str, Any]], document_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        return classify_chunks(chunks, document_metadata)

    def generateSummary(self, text: str) -> str:
        return summarize_text(text)

    def buildEmbeddings(self, document_metadata: Dict[str, Any], chunks: List[Dict[str, Any]]) -> Dict[str, List[Any]]:
        return {
            "ids": [chunk["vector_id"] for chunk in chunks],
            "documents": [chunk["content"] for chunk in chunks],
            "metadatas": [build_vector_metadata(document_metadata, chunk) for chunk in chunks],
        }

    def upsertToVectorStore(self, embedding_payload: Dict[str, List[Any]]) -> None:
        collection.upsert(
            ids=embedding_payload["ids"],
            documents=embedding_payload["documents"],
            metadatas=embedding_payload["metadatas"],
        )

    def updateIndexStatuses(self, doc_id: str, **fields: Any) -> None:
        fields["updated_at"] = now_iso()
        self.repository.update_document(doc_id, fields)

    def queue_upload(self, filename: str, content: bytes, content_type: str, provided_title: Optional[str]) -> UploadResponse:
        validate_filename(filename)

        file_size = len(content)
        if file_size > config.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail="File too large. Maximum size is {0}MB".format(config.MAX_FILE_SIZE // (1024 * 1024)),
            )

        title = validate_title(provided_title) if provided_title else default_title_from_filename(filename)
        file_ext = get_file_extension(filename)
        file_format = config.FORMAT_BY_EXT.get(file_ext, "other")
        mime_type = config.guess_mime_type(filename, content_type or "")
        content_hash = sha256_bytes(content)
        document_key = normalize_title_key(title)

        duplicate = self.repository.find_duplicate(document_key, content_hash)
        if duplicate:
            chunks = self.repository.get_chunks_for_doc(str(duplicate["doc_id"]))
            return UploadResponse(
                success=True,
                duplicate=True,
                message="Duplicate upload detected. Existing indexed document was kept.",
                file=to_file_info(duplicate, chunk_count=len(chunks)),
            )

        version = self.repository.get_latest_version(document_key) + 1
        doc_id = str(uuid.uuid4())
        storage_key = self.saveOriginalFile(doc_id, content, file_ext, version)
        pending_metadata = build_pending_document_metadata(
            doc_id=doc_id,
            title=title,
            original_filename=filename,
            file_ext=file_ext,
            mime_type=mime_type,
            file_format=file_format,
            storage_key=storage_key,
            content_hash=content_hash,
            version=version,
        )
        self.repository.insert_document(pending_metadata, file_size)
        stored_document = self.repository.get_document(doc_id) or pending_metadata
        stored_document["file_size"] = file_size

        return UploadResponse(
            success=True,
            message="File uploaded successfully. Automatic analysis started.",
            processing=True,
            file=to_file_info(stored_document, chunk_count=0),
        )

    def process_pending_upload(self, doc_id: str) -> None:
        document = self.repository.get_document(doc_id)
        if not document:
            return

        stage_status = {
            "parser_status": config.STATUS_PENDING,
            "chunk_status": config.STATUS_PENDING,
            "embedding_status": config.STATUS_PENDING,
            "indexing_status": config.STATUS_PENDING,
        }

        try:
            text = self.parseDocument(str(document["storage_key"]), str(document["format"]))
            if not text.strip():
                raise HTTPException(status_code=400, detail="Could not extract text from file")

            document_metadata = self.classifyDocument(
                doc_id=str(document["doc_id"]),
                title=str(document["title"]),
                original_filename=str(document["original_filename"]),
                file_ext=str(document["file_ext"]),
                mime_type=str(document["mime_type"]),
                file_format=str(document["format"]),
                storage_key=str(document["storage_key"]),
                content_hash=str(document["content_hash"]),
                version=int(document["version"]),
                text=text,
            )
            document_metadata["summary"] = self.generateSummary(text)
            document_metadata["chunk_status"] = config.STATUS_PENDING
            document_metadata["embedding_status"] = config.STATUS_PENDING
            document_metadata["indexing_status"] = config.STATUS_PENDING
            document_metadata["error_message"] = None
            document_metadata["updated_at"] = now_iso()
            self.repository.update_document(doc_id, document_metadata)
            stage_status["parser_status"] = config.STATUS_COMPLETED

            chunks = self.chunkDocument(text)
            classified_chunks = self.classifyChunks(chunks, document_metadata)
            self.repository.delete_chunks_for_doc(doc_id)
            self.repository.insert_chunks(classified_chunks)
            self.updateIndexStatuses(doc_id, chunk_status=config.STATUS_COMPLETED)
            stage_status["chunk_status"] = config.STATUS_COMPLETED

            embedding_payload = self.buildEmbeddings(document_metadata, classified_chunks)
            self.upsertToVectorStore(embedding_payload)
            self.updateIndexStatuses(
                doc_id,
                embedding_status=config.STATUS_COMPLETED,
                indexing_status=config.STATUS_COMPLETED,
                last_indexed_at=now_iso(),
                error_message=None,
            )
            stage_status["embedding_status"] = config.STATUS_COMPLETED
            stage_status["indexing_status"] = config.STATUS_COMPLETED
        except Exception as exc:
            error_message = exc.detail if isinstance(exc, HTTPException) else str(exc)
            self.repository.update_document(
                doc_id,
                {
                    "enabled": False,
                    "review_status": "blocked",
                    "parser_status": stage_status["parser_status"] if stage_status["parser_status"] == config.STATUS_COMPLETED else config.STATUS_FAILED,
                    "chunk_status": stage_status["chunk_status"] if stage_status["chunk_status"] == config.STATUS_COMPLETED else config.STATUS_FAILED,
                    "embedding_status": stage_status["embedding_status"] if stage_status["embedding_status"] == config.STATUS_COMPLETED else config.STATUS_FAILED,
                    "indexing_status": stage_status["indexing_status"] if stage_status["indexing_status"] == config.STATUS_COMPLETED else config.STATUS_FAILED,
                    "error_message": error_message,
                    "tags": ["parse_failed"],
                    "updated_at": now_iso(),
                },
            )

    def rebuild_vectors(self, doc_id: str) -> None:
        document = self.repository.get_document(doc_id)
        if not document:
            raise HTTPException(status_code=404, detail="File not found")

        chunks = self.repository.get_chunks_for_doc(doc_id)
        if not chunks:
            raise HTTPException(status_code=400, detail="No chunks available for indexing")

        self.updateIndexStatuses(
            doc_id,
            embedding_status=config.STATUS_PENDING,
            indexing_status=config.STATUS_PENDING,
            error_message=None,
        )
        self.upsertToVectorStore(self.buildEmbeddings(document, chunks))
        self.updateIndexStatuses(
            doc_id,
            embedding_status=config.STATUS_COMPLETED,
            indexing_status=config.STATUS_COMPLETED,
            last_indexed_at=now_iso(),
            error_message=None,
        )

    def queue_reprocess(self, doc_id: str, action: str) -> None:
        if action not in ALLOWED_ACTIONS:
            raise HTTPException(status_code=400, detail="Unsupported action")

        document = self.repository.get_document(doc_id)
        if not document:
            raise HTTPException(status_code=404, detail="File not found")

        if action == "reindex":
            self.repository.update_document(
                doc_id,
                {
                    "embedding_status": config.STATUS_PENDING,
                    "indexing_status": config.STATUS_PENDING,
                    "error_message": None,
                    "updated_at": now_iso(),
                },
            )
            return

        chunk_ids = [str(chunk.get("vector_id")) for chunk in self.repository.get_chunks_for_doc(doc_id) if chunk.get("vector_id")]
        if chunk_ids:
            collection.delete(ids=chunk_ids)

        self.repository.delete_chunks_for_doc(doc_id)
        self.repository.update_document(
            doc_id,
            {
                "enabled": True,
                "review_status": "needs_review",
                "parser_status": config.STATUS_PENDING,
                "chunk_status": config.STATUS_PENDING,
                "embedding_status": config.STATUS_PENDING,
                "indexing_status": config.STATUS_PENDING,
                "error_message": None,
                "updated_at": now_iso(),
            },
        )

    def run_reprocess(self, doc_id: str, action: str) -> None:
        if action == "reindex":
            self.rebuild_vectors(doc_id)
            return
        self.process_pending_upload(doc_id)


def validate_document_updates(payload: Dict[str, Any]) -> Dict[str, Any]:
    updates: Dict[str, Any] = {}

    if "title" in payload:
        updates["title"] = validate_title(str(payload["title"]))
        updates["document_key"] = normalize_title_key(updates["title"])
    if "source_org" in payload:
        updates["source_org"] = str(payload["source_org"]).strip() or "Unknown"
    if "library" in payload:
        library = str(payload["library"]).strip()
        if library not in ALLOWED_DOC_LIBRARIES:
            raise HTTPException(status_code=400, detail="Invalid library")
        updates["library"] = library
    if "audience" in payload:
        audience = str(payload["audience"]).strip()
        if audience not in ALLOWED_DOC_AUDIENCES:
            raise HTTPException(status_code=400, detail="Invalid audience")
        updates["audience"] = audience
    if "age_bands" in payload:
        updates["age_bands"] = normalize_age_bands(payload["age_bands"])
    if "safety_visibility" in payload:
        visibility = str(payload["safety_visibility"]).strip()
        if visibility not in ALLOWED_DOC_VISIBILITIES:
            raise HTTPException(status_code=400, detail="Invalid safety_visibility")
        updates["safety_visibility"] = visibility
    if "topics" in payload:
        updates["topics"] = normalize_list_value(payload["topics"])
    if "summary" in payload:
        updates["summary"] = str(payload["summary"]).strip()
    if "risk_level" in payload:
        risk_level = str(payload["risk_level"]).strip()
        if risk_level not in ALLOWED_RISK_LEVEL:
            raise HTTPException(status_code=400, detail="Invalid risk_level")
        updates["risk_level"] = risk_level
    if "enabled" in payload:
        updates["enabled"] = bool(payload["enabled"])
    if "review_status" in payload:
        review_status = str(payload["review_status"]).strip()
        if review_status not in ALLOWED_REVIEW_STATUS:
            raise HTTPException(status_code=400, detail="Invalid review_status")
        updates["review_status"] = review_status

    return updates


def validate_chunk_updates(payload: Dict[str, Any]) -> Dict[str, Any]:
    updates: Dict[str, Any] = {}
    if "age_bands" in payload:
        updates["age_bands"] = normalize_age_bands(payload["age_bands"])
    if "audience" in payload:
        audience = str(payload["audience"]).strip()
        if audience not in ALLOWED_CHUNK_AUDIENCES:
            raise HTTPException(status_code=400, detail="Invalid chunk audience")
        updates["audience"] = audience
    if "visibility" in payload:
        visibility = str(payload["visibility"]).strip()
        if visibility not in ALLOWED_CHUNK_VISIBILITIES:
            raise HTTPException(status_code=400, detail="Invalid chunk visibility")
        updates["visibility"] = visibility
    if "topics" in payload:
        updates["topics"] = normalize_list_value(payload["topics"])
    if "retrieval_enabled" in payload:
        updates["retrieval_enabled"] = bool(payload["retrieval_enabled"])

    if updates.get("visibility") == "blocked":
        updates["retrieval_enabled"] = False

    return updates


@app.get("/")
async def root():
    return {"status": "ok", "service": "Knowledge Base API", "version": "2.1.0"}


@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat(), "version": "2.1.0"}


@app.post("/api/upload", response_model=UploadResponse)
async def upload_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
):
    pipeline = KnowledgeIngestionPipeline(get_repository())
    content = await file.read()
    response = pipeline.queue_upload(file.filename, content, file.content_type or "", title)
    if response.processing and response.file:
        background_tasks.add_task(pipeline.process_pending_upload, response.file.id)
    return response


@app.get("/api/files", response_model=FileListResponse)
async def list_files(
    search: Optional[str] = Query(None, description="Search in document title/filename and chunk content"),
    audience: Optional[str] = Query(None, description="Filter by document audience"),
    review_status: Optional[str] = Query(None, description="Filter by review status"),
):
    repository = get_repository()
    documents = repository.list_documents(search=search, audience=audience, review_status=review_status)
    files = [to_file_info(item, chunk_count=int(item.get("chunk_count") or 0)) for item in documents]
    return FileListResponse(success=True, files=files, total=len(files))


@app.get("/api/files/{file_id}")
async def get_file_detail(file_id: str):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    chunks = repository.get_chunks_for_doc(file_id)
    return {
        "success": True,
        "file": to_file_info(document, chunk_count=len(chunks)).model_dump(),
        "document": {
            **document,
            "file_size": int(document.get("file_size") or 0),
            "chunk_count": len(chunks),
        },
        "stages": build_processing_stages(document),
        "chunk_stats": {
            "total": len(chunks),
            "retrieval_enabled": len([chunk for chunk in chunks if chunk.get("retrieval_enabled")]),
            "needs_review": len([chunk for chunk in chunks if float(chunk.get("confidence") or 0.0) < 0.65]),
        },
    }


@app.get("/api/files/{file_id}/status")
async def get_file_status(file_id: str):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    chunk_count = len(repository.get_chunks_for_doc(file_id))
    return get_document_status_payload(document, chunk_count=chunk_count)


@app.put("/api/files/{file_id}")
async def update_file(file_id: str, payload: Dict[str, Any] = Body(...)):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    updates = validate_document_updates(payload)
    if not updates:
        raise HTTPException(status_code=400, detail="No supported fields to update")

    updates["updated_at"] = now_iso()
    repository.update_document(file_id, updates)

    if set(updates.keys()) & {"title", "source_org", "library", "enabled", "review_status"}:
        sync_document_vectors(file_id)

    updated = repository.get_document(file_id) or document
    chunk_count = len(repository.get_chunks_for_doc(file_id))
    return {
        "success": True,
        "message": "Document updated successfully",
        "file": to_file_info(updated, chunk_count=chunk_count).model_dump(),
        "document": {**updated, "chunk_count": chunk_count},
    }


@app.get("/api/files/{file_id}/chunks")
async def list_file_chunks(
    file_id: str,
    search: Optional[str] = Query(None),
    visibility: Optional[str] = Query(None),
    audience: Optional[str] = Query(None),
    age_band: Optional[str] = Query(None),
    retrieval_enabled: Optional[int] = Query(None),
    sort_by: str = Query("chunk_index"),
    sort_dir: str = Query("asc"),
):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    chunks = repository.get_chunks_for_doc(file_id)

    if search:
        needle = search.lower().strip()
        chunks = [
            chunk for chunk in chunks
            if needle in str(chunk.get("content") or "").lower()
            or needle in str(chunk.get("heading_path") or "").lower()
            or needle in " ".join(chunk.get("topics") or []).lower()
        ]
    if visibility:
        chunks = [chunk for chunk in chunks if str(chunk.get("visibility")) == visibility]
    if audience:
        chunks = [chunk for chunk in chunks if str(chunk.get("audience")) == audience]
    if age_band:
        chunks = [chunk for chunk in chunks if age_band in list(chunk.get("age_bands") or [])]
    if retrieval_enabled is not None:
        flag = bool(int(retrieval_enabled))
        chunks = [chunk for chunk in chunks if bool(chunk.get("retrieval_enabled")) == flag]

    reverse = sort_dir.lower() == "desc"
    if sort_by == "confidence":
        chunks = sorted(chunks, key=lambda item: float(item.get("confidence") or 0.0), reverse=reverse)
    else:
        chunks = sorted(chunks, key=lambda item: int(item.get("chunk_index") or 0), reverse=reverse)

    payload = [to_chunk_payload(chunk) for chunk in chunks]
    return {"success": True, "chunks": payload, "total": len(payload)}


@app.get("/api/files/{file_id}/chunks/{chunk_id}")
async def get_file_chunk(file_id: str, chunk_id: str):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    chunk = repository.get_chunk(chunk_id)
    if not chunk or str(chunk.get("doc_id")) != file_id:
        raise HTTPException(status_code=404, detail="Chunk not found")

    return {"success": True, "chunk": to_chunk_payload(chunk)}


@app.put("/api/files/{file_id}/chunks/bulk")
async def bulk_update_file_chunks(file_id: str, payload: Dict[str, Any] = Body(...)):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    chunk_ids = payload.get("chunk_ids") or []
    if not isinstance(chunk_ids, list) or not chunk_ids:
        raise HTTPException(status_code=400, detail="chunk_ids is required")

    chunk_ids = [str(item).strip() for item in chunk_ids if str(item).strip()]
    fields = validate_chunk_updates(payload.get("fields") or {})
    if not fields:
        raise HTTPException(status_code=400, detail="No supported chunk fields to update")

    fields["updated_at"] = now_iso()
    affected = repository.update_chunks(file_id, chunk_ids, fields)
    sync_document_vectors(file_id, chunk_ids=chunk_ids)

    chunks = [repository.get_chunk(chunk_id) for chunk_id in chunk_ids]
    chunks = [to_chunk_payload(chunk) for chunk in chunks if chunk]
    return {"success": True, "updated": affected, "chunks": chunks}


@app.post("/api/files/{file_id}/actions/{action}")
async def queue_file_action(file_id: str, action: str, background_tasks: BackgroundTasks):
    pipeline = KnowledgeIngestionPipeline(get_repository())
    pipeline.queue_reprocess(file_id, action)
    background_tasks.add_task(pipeline.run_reprocess, file_id, action)
    document = get_repository().get_document(file_id)
    chunk_count = len(get_repository().get_chunks_for_doc(file_id))
    return {
        "success": True,
        "message": "Action queued successfully",
        "action": action,
        "status": get_document_status_payload(document or {}, chunk_count=chunk_count),
    }


@app.delete("/api/files/{file_id}", response_model=DeleteResponse)
async def delete_file(file_id: str):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    chunks = repository.get_chunks_for_doc(file_id)
    chunk_ids = [chunk["vector_id"] for chunk in chunks if chunk.get("vector_id")]
    if chunk_ids:
        collection.delete(ids=chunk_ids)

    storage_key = str(document.get("storage_key") or "")
    file_path = get_storage_path(storage_key)
    if storage_key and os.path.exists(file_path):
        os.remove(file_path)

    repository.delete_document(file_id)
    return DeleteResponse(success=True, message="File deleted successfully")


@app.put("/api/files/{file_id}/rename", response_model=UploadResponse)
async def rename_file(file_id: str, new_name: str = Query(..., description="New title")):
    repository = get_repository()
    document = repository.get_document(file_id)
    if not document:
        raise HTTPException(status_code=404, detail="File not found")

    title = validate_title(new_name)
    repository.update_document(
        file_id,
        {
            "title": title,
            "document_key": normalize_title_key(title),
            "updated_at": now_iso(),
        },
    )
    update_vector_titles(file_id, title)

    updated = repository.get_document(file_id) or document
    chunks = repository.get_chunks_for_doc(file_id)
    return UploadResponse(
        success=True,
        message="Title updated successfully",
        file=to_file_info(updated, chunk_count=len(chunks)),
    )


@app.get("/api/search", response_model=SearchResponse)
async def search_knowledge(
    query: str = Query(..., description="Search query"),
    limit: int = Query(5, description="Maximum number of results"),
    session_type: str = Query("system", description="child, parent, or system"),
    age_band: Optional[str] = Query(None, description="0_3, 3_6, 6_12, 12_18"),
    include_filtered: bool = Query(False, description="Return filtered-out chunks with reasons"),
):
    if not query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    if session_type not in {"child", "parent", "system"}:
        raise HTTPException(status_code=400, detail="Invalid session_type")

    indexed_document_count = get_repository().count_indexed_documents()
    if indexed_document_count == 0:
        return SearchResponse(
            success=True,
            results=[],
            filtered_out=[],
            reliable=False,
            message="当前没有任何已索引文档，无法执行真实检索",
            no_result_reason="no_indexed_documents",
            query_topics=map_query_topics(query),
            topic_filter_applied=bool(map_query_topics(query)),
            missing_topic_content=False,
            query=query,
            session_type=session_type,
            age_band=age_band,
            indexed_document_count=0,
        )

    if include_filtered:
        diagnostics = query_collection_with_diagnostics(query, session_type, age_band, min(limit, 20))
        return SearchResponse(
            success=True,
            results=diagnostics["results"],
            filtered_out=diagnostics["filtered_out"],
            reliable=diagnostics["reliable"],
            message=diagnostics["message"],
            no_result_reason=diagnostics["no_result_reason"],
            query_topics=diagnostics["query_topics"],
            topic_filter_applied=diagnostics["topic_filter_applied"],
            missing_topic_content=diagnostics["missing_topic_content"],
            query=query,
            session_type=session_type,
            age_band=age_band,
            indexed_document_count=indexed_document_count,
        )

    results = query_collection(query, session_type, age_band, min(limit, 20))
    return SearchResponse(
        success=True,
        results=results,
        filtered_out=[],
        reliable=bool(results),
        message="已返回通过相关性阈值的真实检索结果" if results else "当前知识库中没有找到可靠的匹配结果",
        no_result_reason=None if results else "no_reliable_match",
        query_topics=map_query_topics(query),
        topic_filter_applied=bool(map_query_topics(query)),
        missing_topic_content=False,
        query=query,
        session_type=session_type,
        age_band=age_band,
        indexed_document_count=indexed_document_count,
    )


@app.get("/api/context")
async def get_context(
    query: str = Query(..., description="Query to get context for"),
    limit: int = Query(3, description="Number of context chunks"),
    session_type: str = Query("child", description="child, parent, or system"),
    age_band: Optional[str] = Query(None, description="0_3, 3_6, 6_12, 12_18"),
):
    if not query.strip():
        return {"context": "", "sources": []}
    if session_type not in {"child", "parent", "system"}:
        return {"context": "", "sources": []}

    results = query_collection(query, session_type, age_band if session_type == "child" else age_band, limit)
    context_parts = [item.document for item in results]
    sources = []

    for item in results:
        title = str(item.metadata.get("title") or "").strip()
        original_filename = str(item.metadata.get("original_filename") or "").strip()
        sources.append(title or original_filename or "Unknown")

    return {"context": "\n\n".join(context_parts), "sources": list(dict.fromkeys(sources))}


if __name__ == "__main__":
    import uvicorn

    print("Starting Knowledge Base Service on http://{0}:{1}".format(config.HOST, config.PORT))
    print("Upload directory: {0}".format(config.UPLOAD_DIR))
    print("ChromaDB directory: {0}".format(config.CHROMA_PERSIST_DIR))
    uvicorn.run(app, host=config.HOST, port=config.PORT)
